from pathlib import Path
from zipfile import ZipFile
from lxml import etree
from docx import Document


DOCX = Path(__file__).resolve().parent / "VulnFlanker_微信公众号推文.docx"
NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "pr": "http://schemas.openxmlformats.org/package/2006/relationships",
}


def require(condition, message):
    if not condition:
        raise AssertionError(message)


doc = Document(DOCX)
require(len(doc.sections) == 1, "expected one section")
sec = doc.sections[0]
require(round(sec.page_width.inches, 2) == 8.50, "page width must be Letter")
require(round(sec.page_height.inches, 2) == 11.00, "page height must be Letter")
for name, value in {
    "top": sec.top_margin.inches,
    "right": sec.right_margin.inches,
    "bottom": sec.bottom_margin.inches,
    "left": sec.left_margin.inches,
}.items():
    require(round(value, 2) == 1.00, f"{name} margin must be 1 inch")
require(round(sec.header_distance.inches, 3) == 0.492, "header distance mismatch")
require(round(sec.footer_distance.inches, 3) == 0.492, "footer distance mismatch")

paragraphs = doc.paragraphs
full_text = "\n".join(p.text for p in paragraphs)
require(len(full_text) > 3500, "article text unexpectedly short")
for token in ["TODO", "TBD", "{{", "}}", "[[", "]]", "PLACEHOLDER"]:
    require(token not in full_text, f"placeholder token remains: {token}")
require("VulnFlanker：" in full_text, "title missing")
require("https://github.com/ZonWin/VulnFlanker" not in full_text, "raw URL should be a hyperlink, not body text")

h1 = [p.text for p in paragraphs if p.style.name == "Heading 1"]
h2 = [p.text for p in paragraphs if p.style.name == "Heading 2"]
require(len(h1) >= 7, "too few Heading 1 sections")
require(len(h2) >= 8, "too few Heading 2 sections")
require(len([p for p in paragraphs if p.style.name == "VF Caption"]) == 4, "expected four captions")
require(len([p for p in paragraphs if p.style.name == "VF Code"]) == 2, "expected two code blocks")
require(len(doc.tables) == 0, "prose should not be packaged in tables")

with ZipFile(DOCX) as zf:
    document_xml = etree.fromstring(zf.read("word/document.xml"))
    styles_xml = etree.fromstring(zf.read("word/styles.xml"))
    numbering_xml = etree.fromstring(zf.read("word/numbering.xml"))
    rels_xml = etree.fromstring(zf.read("word/_rels/document.xml.rels"))
    footer_xml = etree.fromstring(zf.read("word/footer1.xml"))

inline = document_xml.xpath("//wp:inline", namespaces=NS)
anchors = document_xml.xpath("//wp:anchor", namespaces=NS)
require(len(inline) == 4, "expected four inline images")
require(len(anchors) == 0, "floating images are not allowed")
for item in inline:
    extent = item.xpath("./wp:extent", namespaces=NS)[0]
    width_in = int(extent.get("cx")) / 914400
    height_in = int(extent.get("cy")) / 914400
    require(abs(width_in - 6.45) < 0.01, "image width mismatch")
    require(3.20 < height_in < 3.35, "image aspect/height mismatch")
    doc_pr = item.xpath("./wp:docPr", namespaces=NS)[0]
    require(bool(doc_pr.get("descr")), "image alt text missing")

num_paras = document_xml.xpath("//w:p[w:pPr/w:numPr]", namespaces=NS)
require(len(num_paras) >= 17, "real numbered/bulleted paragraphs missing")
abstracts = numbering_xml.xpath("//w:abstractNum", namespaces=NS)
formats = numbering_xml.xpath("//w:numFmt/@w:val", namespaces=NS)
require("bullet" in formats and "decimal" in formats, "bullet/decimal definitions missing")
require(len(abstracts) >= 2, "numbering definitions missing")

hyperlinks = [r for r in rels_xml.xpath("//pr:Relationship", namespaces=NS) if r.get("Type", "").endswith("/hyperlink")]
require(len(hyperlinks) == 1, "expected one external hyperlink")
require(hyperlinks[0].get("Target") == "https://github.com/ZonWin/VulnFlanker", "GitHub hyperlink mismatch")

page_fields = footer_xml.xpath("//w:instrText[contains(., 'PAGE')]", namespaces=NS)
require(len(page_fields) >= 1, "page number field missing")

fake_bullets = [
    p.text
    for p in paragraphs
    if not p.style.name.startswith("Heading")
    and p.text.lstrip().startswith(("•", "- ", "1. ", "2. "))
]
require(not fake_bullets, "fake bullet/numbered paragraphs detected")

normal = styles_xml.xpath("//w:style[@w:styleId='Normal']", namespaces=NS)[0]
normal_east_asia = normal.xpath("./w:rPr/w:rFonts/@w:eastAsia", namespaces=NS)
require(normal_east_asia == ["Microsoft YaHei"], "Normal CJK font mismatch")

print("STRUCTURAL_QA=PASS")
print(f"paragraphs={len(paragraphs)}")
print(f"text_chars={len(full_text)}")
print(f"heading1={len(h1)} heading2={len(h2)}")
print(f"inline_images={len(inline)} captions=4")
print(f"numbered_or_bulleted_paragraphs={len(num_paras)}")
print("page=Letter margins=1in header_footer=0.492in")
