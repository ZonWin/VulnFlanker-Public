from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets"
OUTPUT = ROOT / "VulnFlanker_微信公众号推文.docx"


# narrative_proposal preset, with two named overrides:
# 1) CJK typography override: Microsoft YaHei for reliable Chinese rendering.
# 2) VulnFlanker brand accent override: #2563EB used only for kicker/callout accents.
FONT_LATIN = "Arial"
FONT_CJK = "Microsoft YaHei"
FONT_MONO = "Consolas"
NAVY = "0B2545"
HEADING_BLUE = "2E74B5"
HEADING_DARK = "1F4D78"
BRAND_BLUE = "2563EB"
INK = "182230"
MUTED = "667085"
LIGHT_BLUE = "EFF6FF"
LIGHT_GRAY = "F4F6F9"
BORDER = "D7E2F0"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_run_font(run, size=None, bold=None, italic=None, color=INK, latin=FONT_LATIN, cjk=FONT_CJK):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), cjk)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_border(paragraph, side="left", color=BRAND_BLUE, size="18", space="8"):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), size)
    edge.set(qn("w:space"), space)
    edge.set(qn("w:color"), color)
    pbdr.append(edge)


def set_paragraph_shading(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_text, fld_end])
    set_run_font(run, size=8.5, color=MUTED)


def add_hyperlink(paragraph, text, url, color=BRAND_BLUE, underline=True):
    part = paragraph.part
    rid = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), FONT_LATIN)
    rfonts.set(qn("w:eastAsia"), FONT_CJK)
    rpr.append(rfonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rpr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rpr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")
    rpr.append(sz)
    new_run.append(rpr)
    text_elm = OxmlElement("w:t")
    text_elm.text = text
    new_run.append(text_elm)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def create_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(e.get(qn("w:abstractNumId"))) for e in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    next_abs = max(existing_abs, default=-1) + 1
    next_num = max(existing_num, default=0) + 1

    def add_definition(abstract_id, num_id, fmt, text, font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "280")
        ppr.extend([tabs, ind])
        lvl.extend([start, num_fmt, lvl_text, suff, lvl_jc, ppr])
        if font:
            rpr = OxmlElement("w:rPr")
            rfonts = OxmlElement("w:rFonts")
            rfonts.set(qn("w:ascii"), font)
            rfonts.set(qn("w:hAnsi"), font)
            rpr.append(rfonts)
            lvl.append(rpr)
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(abstract_id))
        num.append(abs_id)
        numbering.append(num)

    add_definition(next_abs, next_num, "bullet", "•", "Segoe UI Symbol")
    add_definition(next_abs + 1, next_num + 1, "decimal", "%1.")
    return next_num, next_num + 1


def apply_num(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_elm = OxmlElement("w:numId")
    num_id_elm.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_elm])
    ppr.append(num_pr)


doc = Document()
doc.core_properties.title = "VulnFlanker：把漏洞情报、资产影响与验证处置串成一条线"
doc.core_properties.subject = "VulnFlanker 微信公众号推文"
doc.core_properties.author = ""
doc.core_properties.comments = "公众号发布成稿；截图来自项目 README 演示环境。"

section = doc.sections[0]
section.orientation = WD_ORIENT.PORTRAIT
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)
section.different_first_page_header_footer = True

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT_LATIN
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.333

heading_tokens = {
    "Heading 1": (16, HEADING_BLUE, 18, 10),
    "Heading 2": (13, HEADING_BLUE, 12, 6),
    "Heading 3": (12, HEADING_DARK, 8, 4),
}
for name, (size, color, before, after) in heading_tokens.items():
    style = styles[name]
    style.font.name = FONT_LATIN
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.1
    style.paragraph_format.keep_with_next = True
    style.paragraph_format.keep_together = True

caption_style = styles.add_style("VF Caption", WD_STYLE_TYPE.PARAGRAPH)
caption_style.font.name = FONT_LATIN
caption_style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
caption_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
caption_style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
caption_style.font.size = Pt(8.5)
caption_style.font.color.rgb = RGBColor.from_string(MUTED)
caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption_style.paragraph_format.space_before = Pt(4)
caption_style.paragraph_format.space_after = Pt(10)
caption_style.paragraph_format.line_spacing = 1.15
caption_style.paragraph_format.keep_together = True

code_style = styles.add_style("VF Code", WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = FONT_MONO
code_style._element.rPr.rFonts.set(qn("w:ascii"), FONT_MONO)
code_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_MONO)
code_style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
code_style.font.size = Pt(9)
code_style.font.color.rgb = RGBColor.from_string(NAVY)
code_style.paragraph_format.left_indent = Inches(0.18)
code_style.paragraph_format.right_indent = Inches(0.18)
code_style.paragraph_format.space_before = Pt(5)
code_style.paragraph_format.space_after = Pt(10)
code_style.paragraph_format.line_spacing = 1.2
code_style.paragraph_format.keep_together = True

bullet_num_id, decimal_num_id = create_numbering(doc)

# Running header/footer: quiet furniture for a polished multi-page editorial article.
header_p = section.header.paragraphs[0]
header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
header_p.paragraph_format.space_after = Pt(0)
hr = header_p.add_run("VULNFLANKER  ·  侧卫 / 漏洞监测平台")
set_run_font(hr, size=8.5, bold=True, color=MUTED)
footer_p = section.footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fr = footer_p.add_run("VulnFlanker 公众号推文  ·  ")
set_run_font(fr, size=8.5, color=MUTED)
add_page_number(footer_p)


def add_body(text, bold_prefix=None, italic=False, keep=False):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, bold=True, color=INK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=11, italic=italic, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, italic=italic, color=INK)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.333
    p.paragraph_format.keep_together = keep
    return p


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph()
    apply_num(p, bullet_num_id)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    p.paragraph_format.keep_together = True
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=10.8, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=10.8)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.8)
    return p


def add_step(text, bold_prefix):
    p = doc.add_paragraph()
    apply_num(p, decimal_num_id)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.208
    p.paragraph_format.keep_together = True
    r1 = p.add_run(bold_prefix)
    set_run_font(r1, size=10.8, bold=True, color=HEADING_DARK)
    r2 = p.add_run(text[len(bold_prefix):])
    set_run_font(r2, size=10.8)
    return p


def add_callout(label, text, fill=LIGHT_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.keep_together = True
    set_paragraph_shading(p, fill)
    set_paragraph_border(p, "left", BRAND_BLUE, "20", "8")
    r1 = p.add_run(label)
    set_run_font(r1, size=10.8, bold=True, color=HEADING_DARK)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.8, color=INK)
    return p


def add_figure(filename, caption, alt_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(ASSETS / filename), width=Inches(6.45))
    inline = run._r.xpath(".//wp:inline")
    if inline:
        doc_pr = inline[0].find(qn("wp:docPr"))
        if doc_pr is not None:
            doc_pr.set("descr", alt_text)
            doc_pr.set("title", alt_text)
    cap = doc.add_paragraph(style="VF Caption")
    cr = cap.add_run(caption)
    set_run_font(cr, size=8.5, color=MUTED)
    return p, cap


def add_code(lines):
    p = doc.add_paragraph(style="VF Code")
    set_paragraph_shading(p, LIGHT_GRAY)
    set_paragraph_border(p, "left", HEADING_BLUE, "14", "7")
    for i, line in enumerate(lines):
        r = p.add_run(line)
        set_run_font(r, size=9, color=NAVY, latin=FONT_MONO, cjk=FONT_CJK)
        if i != len(lines) - 1:
            r.add_break()
    return p


# Opening block: editorial_cover pattern, compressed for a WeChat article rather than a standalone cover.
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
r = p.add_run("开源工具推荐  /  INTERNAL SECURITY OPERATIONS")
set_run_font(r, size=9.5, bold=True, color=BRAND_BLUE)

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(8)
title.paragraph_format.line_spacing = 1.05
title.paragraph_format.keep_together = True
r = title.add_run("VulnFlanker：\n把漏洞情报、资产影响与验证处置串成一条线")
set_run_font(r, size=26.5, bold=True, color=NAVY)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(14)
subtitle.paragraph_format.line_spacing = 1.2
r = subtitle.add_run("一套面向内部安全运营的漏洞影响评估与受控验证平台")
set_run_font(r, size=13.2, color=HEADING_DARK)

lead = doc.add_paragraph()
lead.paragraph_format.space_after = Pt(14)
lead.paragraph_format.line_spacing = 1.35
r = lead.add_run(
    "真正消耗安全团队的，往往不是“又出现了一个 CVE”，而是后面那串问题：我们的哪些资产受影响？谁来负责？先处理哪一个？判断依据能否复核？VulnFlanker 想做的，就是把这些问题放回同一条工作流里。"
)
set_run_font(r, size=12.2, color=INK)

add_figure(
    "github-1.png",
    "图 1｜VulnFlanker 总览：风险、资产、漏洞与闭环进展在同一视图中汇总。",
    "VulnFlanker 总览仪表盘，展示风险、资产、漏洞和闭环概览。",
)
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
note.paragraph_format.space_after = Pt(0)
nr = note.add_run("注：本文产品截图来自项目 README 的演示环境，数值和日期仅用于界面展示。")
set_run_font(nr, size=8.2, italic=True, color=MUTED)

doc.add_page_break()

doc.add_heading("漏洞情报很多，真正缺的是“落地判断”", level=1)
add_body(
    "漏洞运营最常见的困境，不是信息太少，而是信息、资产和处置彼此割裂。公告里写着厂商、产品和版本，资产台账里却可能只有主机名、IP 和零散组件；匹配结果出来以后，还要再问业务重要性、网络暴露、责任人、验证证据和处置状态。"
)
add_body("如果这些环节分散在订阅平台、CMDB、脚本、表格和聊天群里，安全团队很容易卡在三个问题上：")
add_bullet("谁受影响：漏洞描述能否和真实资产、组件、版本、操作系统及功能对应起来？", "谁受影响：")
add_bullet("先做什么：风险优先级是否结合了资产关键性、暴露情况与证据，而不是只看 CVSS？", "先做什么：")
add_bullet("怎么闭环：处置、复核、验证证据和审计记录能否留在同一条链路中？", "怎么闭环：")
add_callout(
    "一句话理解 VulnFlanker：",
    "它不是再做一张“漏洞列表”，而是把漏洞情报转化为面向具体资产、具体责任和具体动作的风险队列。",
)

doc.add_heading("VulnFlanker 是什么？", level=1)
add_body(
    "VulnFlanker（侧卫）是一套面向内部安全运营的开源漏洞影响评估与受控验证平台。它把漏洞情报、Linux 主机快照、资产与漏洞匹配、风险排序、只读验证任务和审计日志连接起来，目标是形成“情报进入—资产更新—影响判断—风险评估—验证处置”的连续工作流。"
)
add_body(
    "情报侧可接入 CISA KEV、阿里云 AVD 和内置 WatchVuln 采集器；资产侧由 Go 编写的 Linux Agent 上报主机快照和心跳；平台再依据产品、版本、操作系统、功能与暴露规则完成匹配，并生成带有优先级、风险因素、说明和稳定风险代码的风险条目。"
)

doc.add_heading("第一步：把外部情报整理成可匹配的数据", level=2)
add_body(
    "VulnFlanker 先把不同来源的漏洞信息标准化到统一目录中。除了 CISA KEV、AVD 和 WatchVuln，平台还提供可选的 AI 信息补充：可通过 OpenAI 兼容接口辅助提取结构化漏洞信息，也可使用已针对 KIMI 做适配的联网搜索补全能力。"
)
add_body(
    "这里的 AI 更适合被理解为“减少清洗和补录成本的助手”，而不是替代规则、证据和人工判断。标准化数据源越可靠，后续匹配越稳；AI 的价值是在原始情报质量不理想时，帮助运营人员更快补齐厂商、产品、版本或影响范围等字段。"
)
add_figure(
    "github-4.png",
    "图 2｜漏洞情报视图：统一查看 KEV、高危统计、检索字段与 AI 建议处理状态。",
    "VulnFlanker 漏洞情报列表，展示 KEV、高危统计和 AI 建议状态。",
)

doc.add_heading("第二步：让资产不只是一串 IP", level=2)
add_body(
    "只知道“某台主机装了什么”还不够。安全运营还需要知道它属于哪个业务系统、谁负责、由哪个团队承接、是否公网暴露、业务关键性如何。VulnFlanker 的资产侧把这些运营字段和 Agent 上报的系统、组件、网络暴露、在线状态、快照新鲜度放到一起。"
)
add_body(
    "这一步看似是台账建设，实际上决定了风险能否被正确分派：同一个漏洞落在测试机和核心生产系统上，优先级显然不该相同；同一条风险如果找不到责任人，也很难真正闭环。"
)
add_figure(
    "github-3.png",
    "图 3｜资产管理视图：资产、业务系统、责任人、责任团队、关键性和暴露类型统一管理。",
    "VulnFlanker 资产管理页面，展示业务归属、责任人、关键性和暴露类型。",
)

doc.add_heading("第三步：把“可能受影响”变成可解释的风险队列", level=2)
add_body(
    "VulnFlanker 的匹配逻辑不仅看漏洞编号，还会综合产品、版本、操作系统、功能和暴露规则。平台的风险匹配流水线支持权重调整，输出风险分值、优先级、匹配状态、处置状态与验证入口。"
)
add_body(
    "更重要的是，它保留风险因素和判断说明。运营人员看到的不只是一个红色标签，还能沿着风险条目回看：为什么命中、对应哪项资产、证据是否充分、当前由谁处置、是否需要进一步验证。"
)
add_figure(
    "github-2.png",
    "图 4｜风险队列视图：风险、漏洞、资产、匹配状态、处置状态与验证入口集中呈现。",
    "VulnFlanker 风险队列，展示优先级、分值、漏洞、资产、匹配、处置和验证入口。",
)

doc.add_heading("为什么值得推荐？", level=1)

doc.add_heading("1. 它解决的是运营断点，不是数据展示", level=2)
add_body(
    "从情报采集到资产影响，从风险排序到验证证据，VulnFlanker 用统一对象和状态把链路串起来。对安全团队来说，这意味着少一些在多个系统之间复制字段、对齐口径和追问责任人的工作。"
)

doc.add_heading("2. 匹配可解释，优先级可调", level=2)
add_body(
    "企业环境差异很大：有的组织更看重公网暴露，有的组织更看重业务关键性，有的组织需要对 KEV 或特定产品线加权。可调权重和明确风险因素，让平台更容易贴合本地管理规则，也便于复盘“为什么先处理它”。"
)

doc.add_heading("3. 验证强调只读与证据留存", level=2)
add_body(
    "平台可以创建只读验证任务，由 Agent 拉取任务并回传证据。这个边界很重要：当前设计有意不包含自动修复和侵入式 PoC 执行，更适合内部受控环境中的确认、复核和审计，而不是把生产主机变成攻击试验场。"
)

doc.add_heading("4. AI 是可选增强，而不是系统前提", level=2)
add_body(
    "即使不接入外部大模型，情报、资产、规则匹配、风险队列和验证链路仍然成立。需要时再启用 AI 补全，可以把成本、数据边界和自动化程度控制在团队能够接受的范围内。"
)

doc.add_heading("5. 开源、可自建，也把安全边界说清楚", level=2)
add_body(
    "项目采用 Apache License 2.0，提供 React 控制台、Python 后端与 Worker、PostgreSQL、Redis，以及 Go 编写的 Linux Agent。README 同时明确提醒：平台会收集敏感资产信息，适合本地演示、内部使用和小型受控环境，不建议直接暴露在互联网。"
)

doc.add_heading("一条完整链路，五步跑起来", level=1)
add_step("采集并标准化情报：从 CISA KEV、AVD、WatchVuln 等来源接收漏洞信息，整理为统一目录。", "采集并标准化情报：")
add_step("接入资产快照：Linux Agent 上报系统、组件、网络暴露、状态和快照新鲜度。", "接入资产快照：")
add_step("匹配并评估影响：根据产品、版本、操作系统、功能和暴露规则判断资产是否受影响。", "匹配并评估影响：")
add_step("生成并分派风险：结合权重、资产关键性与匹配证据形成风险队列，进入责任与处置流程。", "生成并分派风险：")
add_step("执行只读验证并留痕：创建验证任务，回收 Agent 证据，保留处置和审计记录。", "执行只读验证并留痕：")
add_callout(
    "这条链路的核心价值：",
    "把“外部世界发生了什么”翻译成“内部哪些资产需要谁在什么时候采取什么动作”。",
)

doc.add_heading("它适合谁？又不适合谁？", level=1)
add_body("如果你的团队符合下面几种情况，VulnFlanker 值得进入试用清单：")
add_bullet("已有漏洞订阅或情报来源，但影响评估仍大量依赖人工查询和表格。")
add_bullet("管理着一批 Linux 服务器，需要把组件、版本、暴露和业务归属放进同一套判断流程。")
add_bullet("希望在内部建立可解释、可分派、可复核的漏洞风险队列。")
add_bullet("需要用只读方式确认风险，并保留验证证据和审计记录。")
add_bullet("希望自建部署，并根据本地管理规则调整匹配权重和 AI 接入方式。")

add_body("同时也要看清它的边界：")
add_bullet("它不是面向互联网的 SaaS 控制台；共享或生产部署应增加 HTTPS、身份认证边界、网络访问控制和监控。")
add_bullet("它不是自动化攻击平台；当前验证任务坚持只读，不提供侵入式 PoC 或自动修复。")
add_bullet("它不能替代高质量资产治理；资产、版本和责任关系越准确，匹配和分派越可靠。")
add_bullet("它也不是“接上大模型就自动正确”；AI 补全仍应接受规则、来源和人工复核。")

doc.add_heading("快速上手：先在本地跑一个演示环境", level=1)
add_body("项目提供面向演示与开发的 Docker Compose 配置。准备好 Docker 和 Docker Compose 后，可按 README 的方式启动：")
add_code([
    "Copy-Item .env.example .env",
    "# 编辑 .env，至少替换 Redis 密码、Webhook Token 等默认值",
    "docker compose --env-file .env -f .\\deploy\\docker-compose.yml up --build -d",
])
add_body("启动后访问控制台：")
add_code(["http://127.0.0.1:8100/"])
add_body(
    "如需接入 Linux 主机，还要先构建对应架构的 Agent 二进制，并确保 Agent 使用可访问的平台地址，而不是远程主机上的 127.0.0.1。生产或共享环境请启用 TLS，轮换示例密码与令牌，并妥善保管 AI 密钥加密材料。"
)

doc.add_heading("写在最后", level=1)
add_body(
    "漏洞治理最难的部分，往往不在“发现”，而在“把外部情报变成内部行动”。VulnFlanker 的可取之处，是没有把重点停留在更多数据和更多告警上，而是继续向资产影响、责任分派、风险排序、只读验证和审计闭环推进。"
)
add_body(
    "它目前更适合内部试用和小型受控环境，也还有可以持续加固和扩展的空间。但如果你正在寻找一套可自建、可解释、能把漏洞情报真正落到资产和处置上的开源工具，VulnFlanker 值得认真看一眼。"
)
add_callout(
    "一句话推荐：",
    "用一套可控、可解释、可审计的链路，把漏洞情报真正落到资产和处置上。",
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(4)
r = p.add_run("项目地址：")
set_run_font(r, size=11, bold=True, color=INK)
add_hyperlink(p, "github.com/ZonWin/VulnFlanker", "https://github.com/ZonWin/VulnFlanker")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
r = p.add_run("开源协议：Apache License 2.0")
set_run_font(r, size=10.5, color=MUTED)

doc.save(OUTPUT)
print(OUTPUT)
